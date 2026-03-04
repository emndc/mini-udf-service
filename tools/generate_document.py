#!/usr/bin/env python3
"""Generate DOCX, UDF and PDF from Anlaşma Belgesi placeholder data.

Endpoints expected by the UI:
    POST /api/generate/docx  →  filled DOCX
    POST /api/generate/udf   →  UDF (zip containing content.xml)
    POST /api/generate/pdf   →  PDF

Each endpoint receives ``{ placeholders: { ... } }`` and returns the binary.
"""
import io
import os
import sys
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.table import Table as DocxTable

# project root
_PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(_PROJECT_ROOT))

from tools.docx_template_filler import fill_template, generate_output_filename

# ── UDF constants ───────────────────────────────────────────────────────
UDF_TEMPLATE = '''<?xml version="1.0" encoding="UTF-8" ?>
<template format_id="1.8">
<content><![CDATA[{content}]]></content>
{properties}
<elements resolver="hvl-default">
{elements}
</elements>
<styles><style name="default" description="Geçerli" family="Dialog" size="12" bold="false" italic="false" foreground="-13421773" FONT_ATTRIBUTE_KEY="javax.swing.plaf.FontUIResource[family=Dialog,name=Dialog,style=plain,size=12]" /><style name="hvl-default" family="Tahoma" size="12" description="Gövde" /></styles>
</template>'''

DEFAULT_PROPERTIES_XML = (
    '<properties><pageFormat mediaSizeName="1" '
    'leftMargin="42.525000000000006" rightMargin="42.525000000000006" '
    'topMargin="42.525000000000006" bottomMargin="42.52500000000006" '
    'paperOrientation="1" headerFOffset="20.0" footerFOffset="20.0" />'
    '</properties>'
)

HEADER_SPACER_WIDTH = 204.0
HEADER_SPACER_HEIGHT = 60.0
FOOTER_LINE_B64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAgIAAAABCAYAAACsTsZLAAAAGklEQVR42mMICQn5"
    "P4pH8SgexaN4FI/ikYkBsWP6JF+/+nsAAAASUVORK5CYII="
)


# =====================================================================
#  1.  DOCX generation  (template fill)
# =====================================================================
def generate_docx(ui_json: dict,
                  template_name: str = 'AnlasmaBelgesi-#Dolu_v1.docx') -> bytes:
    """Fill the DOCX template with *ui_json* placeholders and return bytes."""
    return fill_template(ui_json, template_name=template_name)


# =====================================================================
#  2.  UDF generation  (DOCX → UDF)
# =====================================================================
def generate_udf(ui_json: dict,
                 template_name: str = 'AnlasmaBelgesi-#Dolu_v1.docx') -> bytes:
    """Fill the template, then convert the resulting DOCX to UDF bytes."""
    docx_bytes = fill_template(ui_json, template_name=template_name)
    return _docx_bytes_to_udf(docx_bytes)


def _docx_bytes_to_udf(docx_bytes: bytes) -> bytes:
    """Convert in-memory DOCX bytes to UDF (ZIP with content.xml).

    Uses the same logic as ``app.py  docx_to_udf_converter`` so that both
    conversion paths produce identical UDF output.
    """
    document = Document(io.BytesIO(docx_bytes))

    content = []
    elements = []
    current_offset = 0

    # ── paragraph helper (identical to app.py) ──────────────────────
    def paragraph_to_element(paragraph: DocxParagraph) -> str:
        nonlocal current_offset
        para_text = paragraph.text or '\u200B'
        para_text = para_text.replace('\t', '    ')

        alignment = "0"
        if paragraph.alignment is not None:
            if paragraph.alignment == 1:
                alignment = "1"
            elif paragraph.alignment == 2:
                alignment = "2"
            elif paragraph.alignment == 3:
                alignment = "3"
            elif paragraph.alignment == 0:
                alignment = "0"

        para_element = f'<paragraph Alignment="{alignment}">'

        # Collect effective text from runs (keeps CDATA in sync with offsets)
        run_text_pieces = []
        for run in paragraph.runs:
            run_text = run.text or ''
            if not run_text:
                continue

            run_text_converted = run_text.replace('\t', '    ')
            run_text_pieces.append(run_text_converted)
            font_size = "12"
            if run.font.size:
                font_size = str(int(run.font.size.pt))
            font_family = run.font.name if run.font.name else "Tahoma"

            style_attrs = [
                f'family="{font_family}"',
                f'size="{font_size}"',
                f'startOffset="{current_offset}"',
                f'length="{len(run_text_converted)}"'
            ]

            if run.bold is True:
                style_attrs.append('bold="true"')
            if run.italic is True:
                style_attrs.append('italic="true"')
            if run.underline is True:
                style_attrs.append('underline="true"')

            para_element += f'<content {" ".join(style_attrs)} />'
            current_offset += len(run_text_converted)

        if run_text_pieces:
            # Use actual run texts for CDATA (guarantees offset alignment)
            content.append(''.join(run_text_pieces))
        else:
            # No effective run text — use paragraph text with fallback content element
            content.append(para_text)
            para_element += f'<content startOffset="{current_offset}" length="{len(para_text)}" family="Tahoma" size="12" />'
            current_offset += len(para_text)

        para_element += '</paragraph>'
        return para_element

    # ── table helpers (identical to app.py — uses xpath) ────────────
    def calc_col_spans(tbl):
        grid_cols = tbl._tbl.xpath('.//w:tblGrid/w:gridCol')
        widths = []
        for col in grid_cols:
            w = col.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
            if w:
                try:
                    widths.append(int(w))
                except ValueError:
                    pass

        if widths:
            total = sum(widths)
            spans = []
            running = 0
            for idx, w in enumerate(widths):
                if total == 0:
                    val = 0
                elif idx == len(widths) - 1:
                    val = max(1, 100 - running)
                else:
                    val = max(1, round(100 * w / total))
                    running += val
                spans.append(val)
            return ','.join(str(v) for v in spans)

        # Fallback: equal distribution if grid widths are missing
        col_count = len(tbl.columns) if hasattr(tbl, 'columns') else 1
        if col_count > 0:
            base_width = 100 // col_count
            col_widths = [base_width] * col_count
            col_widths[-1] = 100 - sum(col_widths[:-1])
            return ','.join(str(v) for v in col_widths)
        return '100'

    def _estimate_row_span(row, col_span_list, col_count):
        """Return a rowSpan value proportional to content height."""
        BASE = 510          # single-line row height
        LINE_H = 340        # additional height per extra line
        USABLE_W_PT = 510.0 # A4 usable width in pt (approx)
        AVG_CHAR_W = 6.0    # average char width in pt at ~10pt font
        max_lines = 1
        for c_idx, cell in enumerate(row.cells):
            pct = col_span_list[c_idx] if c_idx < len(col_span_list) else (100 // max(col_count, 1))
            col_w_pt = USABLE_W_PT * pct / 100.0
            chars_per_line = max(1, int(col_w_pt / AVG_CHAR_W))
            cell_lines = 0
            for para in cell.paragraphs:
                txt = para.text or ''
                if not txt.strip():
                    cell_lines += 1
                else:
                    cell_lines += max(1, -(-len(txt) // chars_per_line))
            max_lines = max(max_lines, cell_lines)
        return max(BASE, max_lines * LINE_H)

    def table_to_element(tbl: DocxTable) -> str:
        nonlocal current_offset
        col_count = len(tbl.columns) if hasattr(tbl, 'columns') else 1
        row_count = len(tbl.rows)

        col_spans = calc_col_spans(tbl)
        col_span_list = [int(v) for v in col_spans.split(',') if v.strip()]
        row_span_values = []
        for row in tbl.rows:
            row_span_values.append(str(_estimate_row_span(row, col_span_list, col_count)))
        row_spans = ','.join(row_span_values) + ','

        table_element_parts = [
            f'<table tableName="Sabit" columnCount="{col_count}" border="borderNone" borderSpec="31" borderColor="-16777216" borderStyle="borderStyle-plain" borderWidth="1.0" columnSpans="{col_spans}" rowSpans="{row_spans}">'  # noqa: E501
        ]

        for r_idx, row in enumerate(tbl.rows, start=1):
            table_element_parts.append(
                f'<row rowName="row{r_idx}" rowType="dataRow" border="borderNone" borderWidth="0.5" height_min="0.8">'
            )
            for cell in row.cells:
                table_element_parts.append('<cell border="borderNone" borderWidth="0.5">')
                for para in cell.paragraphs:
                    table_element_parts.append(paragraph_to_element(para))
                table_element_parts.append('</cell>')
            table_element_parts.append('</row>')

        table_element_parts.append('</table>')
        return ''.join(table_element_parts)

    # ── header / footer helper (identical to app.py) ────────────────
    def build_section_block(tag_name: str) -> str | None:
        collected: list[str] = []

        def add_header_spacer() -> None:
            nonlocal current_offset
            content.append('\u200B')
            spacer = (
                f'<paragraph Alignment="1">'
                f'<content width="{HEADER_SPACER_WIDTH:.1f}" height="{HEADER_SPACER_HEIGHT:.1f}" '
                f'startOffset="{current_offset}" length="1" family="Tahoma" size="12" />'
                f'</paragraph>'
            )
            current_offset += 1
            collected.append(spacer)

        # Only use the first section to avoid duplicate headers/footers
        for section in document.sections:
            part = getattr(section, tag_name, None)
            if not part:
                continue

            # Check if header/footer has actual content (not just empty paras)
            has_content = any(p.text.strip() for p in part.paragraphs) or bool(part.tables)
            if not has_content:
                continue

            if tag_name == 'header':
                add_header_spacer()

            for para in part.paragraphs:
                collected.append(paragraph_to_element(para))
            for tbl in part.tables:
                collected.append(table_to_element(tbl))
            break  # only first section with content

        if collected:
            return f"<{tag_name}>{''.join(collected)}</{tag_name}>"
        return None

    def build_footer_line() -> str:
        nonlocal current_offset
        content.append('\u200B')
        line = (
            f'<paragraph Alignment="1" SpaceBelow="0">'
            f'<image imageData="{FOOTER_LINE_B64}" width="514.0" height="1.0" '
            f'startOffset="{current_offset}" length="1" />'
            f'</paragraph>'
        )
        current_offset += 1
        return line

    # ── Assemble (identical to app.py) ──────────────────────────────
    header_block = build_section_block('header')
    if header_block:
        elements.append(header_block)

    body = document.element.body
    for block in body.iterchildren():
        tag = block.tag
        if tag.endswith('}p'):
            paragraph = DocxParagraph(block, document)
            elements.append(paragraph_to_element(paragraph))
        elif tag.endswith('}tbl'):
            tbl = DocxTable(block, document)
            elements.append(table_to_element(tbl))

    footer_block = build_section_block('footer')
    if footer_block:
        elements.append(footer_block)
    elements.append(build_footer_line())

    if not content:
        content.append('\u200B')
        elements.append('<paragraph Alignment="0"><content startOffset="0" length="1" family="Tahoma" size="12" /></paragraph>')

    # DOCX margins don't map 1:1 to UDF margins (UYAP uses its own layout).
    # Always use the standard UYAP defaults.
    properties_xml = DEFAULT_PROPERTIES_XML

    udf_xml = UDF_TEMPLATE.format(
        content=''.join(content),
        properties=properties_xml,
        elements='\n'.join(elements),
    )

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('content.xml', udf_xml)
    return buf.getvalue()


# =====================================================================
#  3.  PDF generation  (DOCX → PDF direct)
# =====================================================================

# ── Persistent Word COM instance for fast repeated conversions ──────
_word_app = None

def _get_word_app():
    """Return a reusable Word COM Automation instance.

    Keeping Word open avoids the ~10-15 s cold-start penalty on every
    PDF conversion.  The instance is created once and reused.
    """
    global _word_app

    # Ensure COM is initialised on this thread (Flask worker threads
    # don't call CoInitialize automatically).
    try:
        import pythoncom
        pythoncom.CoInitialize()
    except Exception:
        pass

    if _word_app is not None:
        # Check the COM object is still alive
        try:
            _word_app.Documents.Count  # lightweight ping
            return _word_app
        except Exception:
            _word_app = None

    try:
        import win32com.client
        word = win32com.client.DispatchEx('Word.Application')
        word.Visible = False
        word.DisplayAlerts = 0  # wdAlertsNone
        _word_app = word
        return _word_app
    except Exception:
        return None


def _convert_docx_to_pdf_word(docx_path: str, pdf_path: str) -> bool:
    """Convert DOCX → PDF using a persistent Word instance (fast)."""
    word = _get_word_app()
    if word is None:
        return False
    try:
        doc = word.Documents.Open(os.path.abspath(docx_path), ReadOnly=True)
        doc.SaveAs2(os.path.abspath(pdf_path), FileFormat=17)  # 17 = wdFormatPDF
        doc.Close(0)  # wdDoNotSaveChanges
        return os.path.exists(pdf_path)
    except Exception:
        return False


def generate_pdf(ui_json: dict,
                 template_name: str = 'AnlasmaBelgesi-#Dolu_v1.docx') -> bytes:
    """Fill the template, then convert the DOCX directly to PDF.

    Uses a *persistent* Word COM instance for speed (~1-2 s instead of
    ~12 s).  Falls back to ``docx2pdf``, then LibreOffice, then UDF→PDF.
    """
    import shutil
    import subprocess

    docx_bytes = fill_template(ui_json, template_name=template_name)

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, 'filled.docx')
        pdf_path = os.path.join(tmpdir, 'filled.pdf')

        with open(docx_path, 'wb') as f:
            f.write(docx_bytes)

        # 1) Try persistent Word COM instance (fastest, ~1-2 s) ---------------
        print('[generate_pdf] Trying Word COM...')
        if _convert_docx_to_pdf_word(docx_path, pdf_path):
            print(f'[generate_pdf] ✓ Word COM succeeded, PDF size={os.path.getsize(pdf_path)} bytes')
            return Path(pdf_path).read_bytes()
        print('[generate_pdf] ✗ Word COM failed')

        # 2) Try docx2pdf (Word via win32com, cold start ~12 s) ---------------
        print('[generate_pdf] Trying docx2pdf...')
        try:
            import pythoncom
            pythoncom.CoInitialize()
            from docx2pdf import convert as docx2pdf_convert
            docx2pdf_convert(docx_path, pdf_path)
            if os.path.exists(pdf_path):
                print(f'[generate_pdf] ✓ docx2pdf succeeded, PDF size={os.path.getsize(pdf_path)} bytes')
                return Path(pdf_path).read_bytes()
        except Exception as exc:
            print(f'[generate_pdf] ✗ docx2pdf failed: {exc}')

        # 3) Try LibreOffice ---------------------------------------------------
        print('[generate_pdf] Trying LibreOffice...')
        soffice = shutil.which('soffice')
        if soffice:
            try:
                subprocess.run(
                    [soffice, '--headless', '--convert-to', 'pdf',
                     docx_path, '--outdir', tmpdir],
                    check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE,
                )
                # LibreOffice may use the original basename
                alt = os.path.join(tmpdir, 'filled.pdf')
                if os.path.exists(alt):
                    print(f'[generate_pdf] ✓ LibreOffice succeeded')
                    return Path(alt).read_bytes()
            except Exception as exc:
                print(f'[generate_pdf] ✗ LibreOffice failed: {exc}')
        else:
            print('[generate_pdf] ✗ LibreOffice not found')

        # 4) Fallback: DOCX → UDF → PDF via UYAPPDFGenerator ------------------
        print('[generate_pdf] Trying UDF→PDF fallback...')
        udf_bytes = _docx_bytes_to_udf(docx_bytes)
        from tools.udf_to_pdf import UYAPPDFGenerator
        from tools.udf_extract_to_json import decode_cdata_bytes_with_meta

        with zipfile.ZipFile(io.BytesIO(udf_bytes), 'r') as zf:
            raw_xml = zf.read('content.xml')
        meta = decode_cdata_bytes_with_meta(raw_xml)
        parsed = {
            'fields': {},
            'confidences': {},
            'warnings': [],
            'validations': [],
            'metadata': meta,
        }
        generator = UYAPPDFGenerator()
        return generator.create_pdf_birebir(parsed, udf_bytes=udf_bytes)


# =====================================================================
#  Filename helpers
# =====================================================================
def output_filename(ui_json: dict, ext: str = 'docx', template_name: str = 'AnlasmaBelgesi') -> str:
    """Build a descriptive filename based on the template being used."""
    return generate_output_filename(ui_json, template_name=template_name, ext=ext)
