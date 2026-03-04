#!/usr/bin/env python3
"""Fill DOCX templates by replacing #placeholder tokens with actual values.

Usage:
    from tools.docx_template_filler import fill_template
    docx_bytes = fill_template(ui_json, template_name="AnlasmaBelgesi-#Dolu_v1.docx")
"""
import copy
import io
import json
import logging
import os
import re
from datetime import datetime
from pathlib import Path

logger = logging.getLogger(__name__)

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

# ── Template directory ──────────────────────────────────────────────────
TEMPLATES_DIR = Path(__file__).resolve().parents[1] / 'uploads'

# Patterns for cleaning up orphaned label text when a placeholder value is empty.
# After replacement, if a value is "", surrounding boilerplate like
# "T.C. Kimlik No: " or " - " should also be removed.
# Each pattern ONLY matches when the label has NO value next to it.
_CLEANUP_PATTERNS = [
    # "X T.C. Kimlik No: $"  →  "X" (label at end = tckn was empty)
    (r'\s+T\.?C\.?\s*Kimlik\s*No\s*:\s*$', ''),
    # "^ T.C. Kimlik No: " at start with no name before it
    (r'^\s*T\.?C\.?\s*Kimlik\s*No\s*:\s*$', ''),
    # "X Mersis No: $"  →  "X" (mersis was empty)
    (r'\s+Mersis\s*No\s*:\s*$', ''),
    (r'^\s*Mersis\s*No\s*:\s*$', ''),
    # " -   " or " - " alone (separator between gsm and eposta, both empty)
    (r'^\s*-\s*$', ''),
    # trailing " - " when eposta is empty: "05551234567 -   "
    (r'\s+-\s+$', ''),
    # leading " - " when gsm is empty: " -   ahmet@mail.com"
    (r'^\s*-\s{2,}', ''),
    # trailing/leading commas from muzakere lists: ", , , "
    (r'(?:,\s*){2,}', ', '),
    # leading/trailing comma+space
    (r'^\s*,\s*', ''),
    (r'\s*,\s*$', ''),
    # trailing comma before static suffix text (empty x-slots cleaned up)
    (r',\s+(?=alacağı\s+olup)', ' '),
    (r',\s+(?=olmadığı\s+konularında)', ' '),
    # "()" leftover when uyusmazlikKonusu is empty
    (r'\s*\(\s*\)', ''),
    # "(ADB Sicil No: )" when sicil is empty
    (r'\(ADB\s+Sicil\s+No\s*:\s*\)', ''),
    # double+ spaces
    (r'  +', ' '),
]


# ── Public API ──────────────────────────────────────────────────────────
def fill_template(ui_json: dict,
                  template_name: str = 'AnlasmaBelgesi-#Dolu_v1.docx',
                  output_path: str | None = None) -> bytes:
    """Replace every ``#placeholder`` in *template_name* with values from
    *ui_json* and return the resulting DOCX as bytes.

    Parameters
    ----------
    ui_json : dict
        The UI payload (same schema the frontend sends).
    template_name : str
        Filename of the template inside ``uploads/``.
    output_path : str | None
        If given, the filled DOCX is also written to this path.

    Returns
    -------
    bytes
        The filled DOCX file content.
    """
    # ── LOG: dump incoming UI JSON to file for debugging ──
    _log_dir = Path(__file__).resolve().parents[1] / 'outputs'
    _log_dir.mkdir(exist_ok=True)
    _log_file = _log_dir / 'ui_json_incoming.json'
    try:
        with open(_log_file, 'w', encoding='utf-8') as _f:
            json.dump(ui_json, _f, ensure_ascii=False, indent=2, default=str)
        logger.info('UI JSON logged to %s', _log_file)
        print(f'[docx_template_filler] UI JSON logged → {_log_file}')
    except Exception as _e:
        logger.warning('Could not log UI JSON: %s', _e)

    replacements = _build_replacements(ui_json)

    # ── LOG: dump resolved replacements ──
    _repl_file = _log_dir / 'replacements_resolved.json'
    try:
        with open(_repl_file, 'w', encoding='utf-8') as _f:
            json.dump(replacements, _f, ensure_ascii=False, indent=2, default=str)
        print(f'[docx_template_filler] Replacements logged → {_repl_file}')
    except Exception:
        pass

    template_path = TEMPLATES_DIR / template_name
    if not template_path.exists():
        raise FileNotFoundError(f'Template not found: {template_path}')

    doc = Document(str(template_path))

    # Replace in body paragraphs
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    # Replace in tables
    for table in doc.tables:
        _replace_in_table(table, replacements)

    # Replace in headers & footers (including their tables)
    for section in doc.sections:
        for part in (section.header, section.footer,
                     section.first_page_header, section.first_page_footer,
                     section.even_page_header, section.even_page_footer):
            if part is None:
                continue
            for para in part.paragraphs:
                _replace_in_paragraph(para, replacements)
            for tbl in part.tables:
                _replace_in_table(tbl, replacements)

    # ── Static label replacements ──
    # If only one başvurucu, remove "(1)" from "BAŞVURUCU (1)"
    # If only one diğer taraf, remove "(1)" from "DİĞER TARAF (1)"
    basvurucular = ui_json.get('basvurucular') or []
    diger_taraflar = ui_json.get('digerTaraflar') or []
    label_map = {}
    if len(basvurucular) <= 1:
        label_map['BAŞVURUCU (1)'] = 'BAŞVURUCU'
    if len(diger_taraflar) <= 1:
        label_map['DİĞER TARAF (1)'] = 'DİĞER TARAF'
    _replace_static_labels(doc, label_map)

    # ── Make "VEKİLİ" and "YETKİLİSİ" italic ──
    _italicize_keywords(doc, {'VEKİLİ', 'YETKİLİSİ'})

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    if output_path:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, 'wb') as f:
            f.write(docx_bytes)

    return docx_bytes


# ── Template → display name mapping ─────────────────────────────────────
_TEMPLATE_DISPLAY_NAMES: dict[str, str] = {
    'AnlasmaBelgesi': 'Anlasma Belgesi',
    'SonTutanak-Anlasma': 'Son Tutanak',
}


def generate_output_filename(ui_json: dict,
                             template_name: str = 'AnlasmaBelgesi',
                             ext: str = 'docx') -> str:
    """Build a descriptive output filename: ``<Display Name>(<dosya no>).<ext>``

    *template_name* is matched against ``_TEMPLATE_DISPLAY_NAMES`` to choose
    the human-readable prefix.

    Searches multiple possible JSON paths for the dosya numarası:
      - dosyaKayit.arabuluculukDosyaNo
      - dosyaKayit.basvuruDosyaNo
      - dosyaBilgileri …
      - Flat keys (ArabuluculukDosyaNo etc.)
    """
    # Determine display name from template_name
    display = 'Anlasma Belgesi'  # default
    for key, name in _TEMPLATE_DISPLAY_NAMES.items():
        if key in template_name:
            display = name
            break

    dosya_no = ''
    # Try multiple paths — prefer arabuluculukDosyaNo over basvuruDosyaNo
    for section_key in ('dosyaKayit', 'dosyaBilgileri'):
        section = ui_json.get(section_key) or {}
        for field_key in ('arabuluculukDosyaNo', 'basvuruDosyaNo'):
            val = (section.get(field_key) or '').strip()
            if val:
                dosya_no = val
                break
        if dosya_no:
            break
    # Flat keys fallback
    if not dosya_no:
        for flat_key in ('ArabuluculukDosyaNo', 'ArabulucuDosyaNo', 'BasvurucuDosyaNo', 'basvuruDosyaNo'):
            val = (ui_json.get(flat_key) or '').strip()
            if val:
                dosya_no = val
                break

    # Sanitise dosya_no for filesystem
    safe_no = re.sub(r'[^\w\-.]', '_', dosya_no) if dosya_no else ''

    if safe_no:
        return f'{display}({safe_no}).{ext}'
    return f'{display}.{ext}'


# ── Mapping: UI JSON → placeholder dict ─────────────────────────────────

def _format_date(val: str) -> str:
    """Ensure date string is in gün.ay.yıl (DD.MM.YYYY) format.

    Accepts: DD.MM.YYYY, YYYY-MM-DD, DD/MM/YYYY and returns DD.MM.YYYY.
    If the value is already DD.MM.YYYY or can't be parsed, return as-is.
    """
    if not val:
        return val
    val = val.strip()
    # Already DD.MM.YYYY?
    m = re.match(r'^(\d{1,2})\.(\d{1,2})\.(\d{4})$', val)
    if m:
        return f'{int(m.group(1)):02d}.{int(m.group(2)):02d}.{m.group(3)}'
    # YYYY-MM-DD
    m = re.match(r'^(\d{4})-(\d{1,2})-(\d{1,2})$', val)
    if m:
        return f'{int(m.group(3)):02d}.{int(m.group(2)):02d}.{m.group(1)}'
    # DD/MM/YYYY
    m = re.match(r'^(\d{1,2})/(\d{1,2})/(\d{4})$', val)
    if m:
        return f'{int(m.group(1)):02d}.{int(m.group(2)):02d}.{m.group(3)}'
    return val


def _build_replacements(ui: dict) -> dict[str, str]:
    """Convert the structured UI JSON into a flat ``{#placeholder: value}``
    dictionary."""
    # ── Handle nested 'placeholders' structure from frontend ──
    # Frontend may send data nested under 'placeholders' key. Extract it.
    if 'placeholders' in ui and isinstance(ui.get('placeholders'), dict):
        placeholders_data = ui.pop('placeholders', {})
        # Merge: ui still contains top-level data, but also add from placeholders
        for key, val in placeholders_data.items():
            if key not in ui:  # Don't override existing top-level keys
                ui[key] = val
    
    r: dict[str, str] = {}

    def _v(obj, key, default=''):
    
    r: dict[str, str] = {}

    def _v(obj, key, default=''):
        """Safely get a string value.  Also handles ``{deger: …}`` dicts
        produced by ``map_to_ui_schema``."""
        if obj is None:
            return default
        val = obj.get(key)
        if val is None or val == '':
            return default
        # Handle {deger: ..., tutanaktaYazsin: ...} wrapper
        if isinstance(val, dict) and 'deger' in val:
            inner = val.get('deger')
            if inner is None or inner == '':
                return default
            return str(inner).strip()
        return str(val).strip()

    # ── Arabulucu ──
    arb = ui.get('arabulucu') or {}
    r['#arb-adiSoyadi'] = _v(arb, 'adiSoyadi')
    r['#arb-sicil'] = _v(arb, 'sicilNo')

    # ── Dosya / Büro ──
    # UI may send dosya info under 'dosyaKayit' or 'dosyaBilgileri'
    dk = ui.get('dosyaKayit') or {}
    db = ui.get('dosyaBilgileri') or {}
    r['#ArabulucuBurosu'] = _v(dk, 'arabuluculukBurosu') or _v(db, 'arabulucuBurosu')
    r['#BasvurucuDosyaNo'] = _v(dk, 'basvuruDosyaNo') or _v(db, 'basvuruDosyaNo')
    r['#ArabulucuDosyaNo'] = _v(dk, 'arabuluculukDosyaNo') or _v(db, 'arabuluculukDosyaNo')
    r['#basvuruTarihi'] = _format_date(_v(dk, 'basvuruTarihi') or _v(db, 'basvuruTarihi'))
    r['#goreviKabulTarihi'] = _format_date(_v(dk, 'goreviKabulTarihi') or _v(db, 'goreviKabulTarihi'))

    # ── Başvurucular (index 1-based) ──
    basvurucular = ui.get('basvurucular') or []

    # Pre-fill ALL possible vekil/basvurucu slots with empty string
    # so unused placeholders are cleared even when arrays are short/empty.
    for idx in range(1, 4):  # up to 3 başvurucu
        prefix = f'#basvurucu-gk-{idx}'
        for suffix in ('-adiSoyadi', '-tckn', '-adres', '-gsm', '-eposta', '-temsilci'):
            r.setdefault(f'{prefix}{suffix}', '')
    for vi in range(1, 4):   # up to 3 vekil
        vprefix = f'#b-vekil{vi}'
        for suffix in ('-adiSoyadi', '-tckn', '-gsm', '-eposta'):
            r.setdefault(f'{vprefix}{suffix}', '')

    for i, bsv in enumerate(basvurucular):
        idx = i + 1
        prefix = f'#basvurucu-gk-{idx}'
        r[f'{prefix}-adiSoyadi'] = _v(bsv, 'adiSoyadi')
        # tckn: flat 'tckn' (test_v2) OR 'tcKimlikNo' (UI schema)
        r[f'{prefix}-tckn'] = _v(bsv, 'tckn') or _v(bsv, 'tcKimlikNo')
        r[f'{prefix}-adres'] = _v(bsv, 'adres')
        # gsm/eposta: flat key OR nested under iletisim.{field}.deger
        bsv_iletisim = bsv.get('iletisim') or {}
        r[f'{prefix}-gsm'] = _v(bsv, 'gsm') or _v(bsv_iletisim, 'gsm')
        r[f'{prefix}-eposta'] = _v(bsv, 'eposta') or _v(bsv_iletisim, 'eposta')
        r[f'{prefix}-temsilci'] = _v(bsv, 'temsilci')

        # Vekiller
        vekiller = bsv.get('vekiller') or []
        for vi, vek in enumerate(vekiller):
            vprefix = f'#b-vekil{vi + 1}'
            r[f'{vprefix}-adiSoyadi'] = _v(vek, 'adiSoyadi')
            r[f'{vprefix}-tckn'] = _v(vek, 'tckn') or _v(vek, 'tcKimlikNo')
            r[f'{vprefix}-gsm'] = _v(vek, 'gsm')
            r[f'{vprefix}-eposta'] = _v(vek, 'eposta')

        # Auto-set temsilci = "VEKİLİ" if vekil exists and temsilci is empty
        if vekiller and any(_v(v, 'adiSoyadi') for v in vekiller):
            if not r.get(f'{prefix}-temsilci'):
                r[f'{prefix}-temsilci'] = 'VEKİLİ'

    # ── Diğer Taraflar (index 1-based) ──
    # Accept both 'digerTaraflar' (array) and 'digerTaraf' (single object)
    diger = ui.get('digerTaraflar') or []
    if not diger:
        dt_single = ui.get('digerTaraf')
        if dt_single and isinstance(dt_single, dict):
            diger = [dt_single]

    # Pre-fill ALL possible digertaraf/yetkili slots
    for idx in range(1, 4):  # up to 3 diğer taraf
        prefix = f'#digertaraf-tk-{idx}'
        for suffix in ('-unvani', '-mersis', '-adres', '-telefon', '-temsilci'):
            r.setdefault(f'{prefix}{suffix}', '')
    for yi in range(1, 4):   # up to 3 yetkili
        yprefix = f'#dt-yetkili{yi}'
        for suffix in ('-adiSoyadi', '-tckn', '-gsm', '-eposta'):
            r.setdefault(f'{yprefix}{suffix}', '')

    for i, dt in enumerate(diger):
        idx = i + 1
        prefix = f'#digertaraf-tk-{idx}'
        # unvani: flat 'unvani' OR 'unvan' (UI schema)
        r[f'{prefix}-unvani'] = _v(dt, 'unvani') or _v(dt, 'unvan')
        # mersisNo may be {deger: ...} dict — _v handles that
        r[f'{prefix}-mersis'] = _v(dt, 'mersisNo') or _v(dt, 'mersis')
        r[f'{prefix}-adres'] = _v(dt, 'adres')
        # telefon: flat OR nested under iletisim
        dt_iletisim = dt.get('iletisim') or {}
        r[f'{prefix}-telefon'] = _v(dt, 'telefon') or _v(dt_iletisim, 'telefon') or _v(dt_iletisim, 'gsm')
        r[f'{prefix}-temsilci'] = _v(dt, 'temsilci')

        # Yetkililer
        yetkililer = dt.get('yetkililer') or []
        for yi, ytk in enumerate(yetkililer):
            yprefix = f'#dt-yetkili{yi + 1}'
            r[f'{yprefix}-adiSoyadi'] = _v(ytk, 'adiSoyadi')
            r[f'{yprefix}-tckn'] = _v(ytk, 'tckn') or _v(ytk, 'tcKimlikNo')
            r[f'{yprefix}-gsm'] = _v(ytk, 'gsm')
            r[f'{yprefix}-eposta'] = _v(ytk, 'eposta')

        # Auto-set temsilci = "YETKİLİSİ" if yetkili exists and temsilci is empty
        if yetkililer and any(_v(y, 'adiSoyadi') for y in yetkililer):
            if not r.get(f'{prefix}-temsilci'):
                r[f'{prefix}-temsilci'] = 'YETKİLİSİ'

        # Vekiller (diğer taraf vekili)
        dt_vekiller = dt.get('vekiller') or []
        for dvi, dvek in enumerate(dt_vekiller):
            dvprefix = f'#dt-vekil{dvi + 1}'
            r[f'{dvprefix}-adiSoyadi'] = _v(dvek, 'adiSoyadi')

    # Pre-fill #dt-vekil slots with empty string
    for dvi in range(1, 4):
        r.setdefault(f'#dt-vekil{dvi}-adiSoyadi', '')

    # ── Uyuşmazlık ──
    uy = ui.get('uyusmazlik') or {}
    r['#uyusmazlikTuru'] = _v(uy, 'uyusmazlikTuru')
    r['#uyusmazlikKonusu'] = _v(uy, 'uyusmazlikKonusu')

    # ── Toplantı Yeri (toplanti section'dan okunur) ──
    tp = ui.get('toplanti') or {}
    r['#toplantiYeri'] = _v(tp, 'toplantiYeri')

    # ── Anlaşma Belgesi ──
    ab = ui.get('anlasmaBelgesi') or {}
    r['#sozlesmeBaslangic'] = _format_date(_v(ab, 'sozlesmeBaslangic'))
    r['#sozlesmeBitis'] = _format_date(_v(ab, 'sozlesmeBitis'))

    # İşten Çıkış Kodu → "İşten Çıkış Kodu: XX"
    cikis_kodu = _v(ab, 'istenCikisKodu')
    r['#istenCikisKodu'] = f'İşten Çıkış Kodu: {cikis_kodu}' if cikis_kodu else ''

    # ── Müzakere hususları ──
    # isciAlacaklari → #muzakere-Husus-01..10 (tutar + parabirimi  – Varılan Anlaşma)
    #                → #muzakere-Husus-x1..x9, x0 (label only – Uyuşmazlık Konusu)
    # Pre-fill ALL slots with empty string so unused ones are cleared.
    for slot in range(1, 11):
        idx = f'{slot:02d}'
        r.setdefault(f'#muzakere-Husus-{idx}', '')
        r.setdefault(f'#muzakere-Husus-{idx}-tutar', '')
        r.setdefault(f'#muzakere-Husus-{idx}-parabirimi', '')
    # x-series: x1..x9, x0  (10 slots for husus labels)
    for xs in list(range(1, 10)) + [0]:  # x1,x2,...,x9,x0
        r.setdefault(f'#muzakere-Husus-x{xs}', '')

    isci = ab.get('isciAlacaklari') or []
    for i, item in enumerate(isci):
        idx = f'{i + 1:02d}'  # 01, 02, ...
        label = _v(item, 'label')
        r[f'#muzakere-Husus-{idx}'] = label
        r[f'#muzakere-Husus-{idx}-tutar'] = _v(item, 'tutar')
        r[f'#muzakere-Husus-{idx}-parabirimi'] = _v(item, 'paraBirimi')
        # Map to x-series (x1 for first item, x2 for second, etc.)
        x_key = i + 1 if i + 1 <= 9 else 0  # x1..x9, overflow → x0
        r[f'#muzakere-Husus-x{x_key}'] = label

    # İşveren alacakları (devam eden indeks)
    isveren = ab.get('isverenAlacaklari') or []
    offset = len(isci)
    for i, item in enumerate(isveren):
        idx = f'{offset + i + 1:02d}'
        label = _v(item, 'label')
        r[f'#muzakere-Husus-{idx}'] = label
        r[f'#muzakere-Husus-{idx}-tutar'] = _v(item, 'tutar')
        r[f'#muzakere-Husus-{idx}-parabirimi'] = _v(item, 'paraBirimi')
        x_key = offset + i + 1 if offset + i + 1 <= 9 else 0
        r[f'#muzakere-Husus-x{x_key}'] = label

    # Net/brüt (genel)
    r['#net-brut'] = _v(ab, 'toplamTutarNetBrut')

    # Toplam tutar
    r['#alacak-toplamtutar'] = _v(ab, 'toplamTutarGirisi') or _v(ab, 'odemeTutari')
    r['#alacak-toplamtutar-netbrut'] = _v(ab, 'toplamTutarNetBrut')
    r['#alacak-toplamtutar-parabirimi'] = _v(ab, 'toplamTutarBirimi') or _v(ab, 'odemeTutariBirimi')

    # Ödeme
    r['#alacak-pesin-sonodeme-tarihi'] = _format_date(_v(ab, 'sonOdemeTarihi'))
    r['#odeme-yapacak-kisi'] = _v(ab, 'odemeYapacakKisi')
    r['#odeme-alacak-kisi'] = _v(ab, 'odemeAlacakKisi')
    r['#bankaadi'] = _v(ab, 'banka')
    r['#alacak-iban-no'] = _v(ab, 'iban')

    # ── Son Tutanak (Anlaşma) ──
    sta = ui.get('sonTutanakAnlasma') or {}
    cozum = _v(sta, 'cozumOnerisi')
    _cozum_map = {
        'bulunulmustur': 'Taraflara çözüm önerisinde bulunulmuştur.',
        'bulunulmamistur': 'Taraflara çözüm önerisinde bulunulmamıştır.',
    }
    r['#cozumonerisi-yes'] = _cozum_map.get(cozum, cozum or '')
    r['#oturum-sayisi'] = _v(sta, 'oturumSayisi')
    r['#oturum-suresi-saat'] = _v(sta, 'oturumSureSaat')
    r['#oturum-suresi-dakika'] = _v(sta, 'oturumSureDakika')

    # ── Belge Meta ──
    bm = ui.get('belgeMeta') or {}
    r['#belge-nusha-sayisi'] = _v(bm, 'nushaSayisi')
    r['#belgenin-tarihi'] = _format_date(_v(bm, 'belgeTarihi'))

    # ── Footer ──
    # Footer lines derived from arabulucu
    arb_name = r.get('#arb-adiSoyadi', '')
    arb_sicil = r.get('#arb-sicil', '')
    r['#arb-footer-line1'] = f'Arabulucu: {arb_name}' if arb_name else ''
    r['#arb-footer-line2'] = f'ADB Sicil No: {arb_sicil}' if arb_sicil else ''

    return r


# ── DOCX replacement helpers ────────────────────────────────────────────

def _replace_static_labels(doc, label_map: dict):
    """Replace static label text (non-placeholder) in document tables.

    Used for things like 'BAŞVURUCU (1)' → 'BAŞVURUCU' when there is
    only one applicant.
    """
    if not label_map:
        return
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    runs = para.runs
                    if not runs:
                        continue
                    full = ''.join(r.text for r in runs)
                    new_full = full
                    for old, new in label_map.items():
                        if old in new_full:
                            new_full = new_full.replace(old, new)
                    if new_full != full:
                        runs[0].text = new_full
                        for r in runs[1:]:
                            r.text = ''
                # Nested tables
                for nested in cell.tables:
                    _replace_static_labels_table(nested, label_map)


def _replace_static_labels_table(table, label_map: dict):
    """Recursive helper for nested tables."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                runs = para.runs
                if not runs:
                    continue
                full = ''.join(r.text for r in runs)
                new_full = full
                for old, new in label_map.items():
                    if old in new_full:
                        new_full = new_full.replace(old, new)
                if new_full != full:
                    runs[0].text = new_full
                    for r in runs[1:]:
                        r.text = ''
            for nested in cell.tables:
                _replace_static_labels_table(nested, label_map)

def _replace_in_table(table, replacements: dict):
    """Replace placeholders in every cell of a table (including nested tables)."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _replace_in_paragraph(para, replacements)
            # Nested tables
            for nested in cell.tables:
                _replace_in_table(nested, replacements)


def _replace_in_paragraph(paragraph, replacements: dict):
    """Replace #placeholder tokens across runs in a paragraph.

    Placeholders may be split across multiple runs by Word (e.g. ``#`` in one
    run, ``ArabulucuBurosu`` in the next).  This method reconstructs the full
    text, performs replacements, and then re-distributes the result back to
    the original runs while preserving formatting.
    After replacement, run colours are reset to black and orphaned label
    fragments are cleaned up.
    """
    runs = paragraph.runs
    if not runs:
        return

    # Build full text from runs
    full_text = ''.join(r.text for r in runs)

    # Quick check: any placeholder in full_text?
    if '#' not in full_text:
        return

    # Perform replacements on the full concatenated text
    # CRITICAL: Sort by key length descending so longer placeholders are
    # replaced first (e.g. #muzakere-Husus-01-parabirimi before #muzakere-Husus-01)
    sorted_keys = sorted(replacements.keys(), key=len, reverse=True)
    new_text = full_text
    for placeholder in sorted_keys:
        value = replacements[placeholder]
        if placeholder in new_text:
            new_text = new_text.replace(placeholder, value)

    # Clean up orphaned labels (e.g. "T.C. Kimlik No: " when value was empty)
    for pattern, repl in _CLEANUP_PATTERNS:
        new_text = re.sub(pattern, repl, new_text)
    new_text = new_text.strip()

    # If nothing changed, skip
    if new_text == full_text:
        return

    # Re-distribute into runs: put all text in first run, clear the rest.
    # This is the simplest approach that preserves the first run's formatting.
    # For better preservation we try to map text back to run boundaries.
    _redistribute_text(runs, full_text, new_text, replacements)

    # ── Post-replacement: reset run colours to black ──
    _reset_run_colors(runs)


def _redistribute_text(runs, old_full, new_full, replacements):
    """Smartly redistribute replaced text back into existing runs.

    Strategy: walk through runs and try to keep non-placeholder text in its
    original run.  When a placeholder spans multiple runs, collapse it into
    the first run that started the placeholder and empty the rest.
    """
    # Build a map of run boundaries  [start, end)
    boundaries = []
    pos = 0
    for r in runs:
        length = len(r.text)
        boundaries.append((pos, pos + length))
        pos += length

    # Find which runs are affected by replacements
    # Simple approach: put entire new text in first run, clear others
    # But try to preserve formatting for the non-placeholder parts first.

    # Check if any placeholder was split across runs
    # Detect by checking if a run starts with part of a placeholder
    has_split = False
    for i, r in enumerate(runs):
        if r.text and not r.text.startswith('#') and '#' in old_full:
            # Check if this run is the continuation of a placeholder from prev run
            start = boundaries[i][0]
            if start > 0:
                # Check text before this run for an unclosed placeholder
                before = old_full[:start]
                last_hash = before.rfind('#')
                if last_hash >= 0:
                    potential = old_full[last_hash:boundaries[i][1]]
                    for ph in replacements:
                        if potential.startswith(ph) or ph.startswith(potential):
                            has_split = True
                            break

    if has_split:
        # Fallback: put all text in first run, clear the rest
        runs[0].text = new_full
        for r in runs[1:]:
            r.text = ''
    else:
        # No cross-run splits: replace within each run independently
        # Sort by length descending (same as above)
        sorted_keys = sorted(replacements.keys(), key=len, reverse=True)
        for r in runs:
            if '#' in r.text:
                text = r.text
                for placeholder in sorted_keys:
                    value = replacements[placeholder]
                    if placeholder in text:
                        text = text.replace(placeholder, value)
                r.text = text


def _italicize_keywords(doc, keywords: set[str]):
    """Find runs containing any of *keywords* and make them italic.

    If a run contains ONLY the keyword, just set italic.
    If the keyword is embedded in a larger text, split it into separate runs
    so only the keyword portion becomes italic.
    """
    from docx.oxml import OxmlElement
    from copy import deepcopy

    def _process_paragraphs(paragraphs):
        for para in paragraphs:
            for run in list(para.runs):
                text = run.text or ''
                if not text:
                    continue
                for kw in keywords:
                    if kw not in text:
                        continue
                    if text.strip() == kw:
                        # Entire run is the keyword — just italicize
                        run.font.italic = True
                    else:
                        # Keyword is embedded — split run into parts
                        idx = text.find(kw)
                        before = text[:idx]
                        after = text[idx + len(kw):]

                        # Keep the 'before' part in the current run
                        run.text = before

                        # Create a new italic run for the keyword
                        kw_run = OxmlElement('w:r')
                        # Copy formatting from original run
                        rpr = run._element.find(qn('w:rPr'))
                        if rpr is not None:
                            kw_run.append(deepcopy(rpr))
                        else:
                            kw_run.append(OxmlElement('w:rPr'))
                        # Set italic
                        kw_rpr = kw_run.find(qn('w:rPr'))
                        i_elem = OxmlElement('w:i')
                        kw_rpr.append(i_elem)
                        # Set text
                        kw_t = OxmlElement('w:t')
                        kw_t.text = kw
                        kw_t.set(qn('xml:space'), 'preserve')
                        kw_run.append(kw_t)

                        # Create a normal run for the 'after' part
                        after_run = OxmlElement('w:r')
                        if rpr is not None:
                            after_run.append(deepcopy(rpr))
                        after_t = OxmlElement('w:t')
                        after_t.text = after
                        after_t.set(qn('xml:space'), 'preserve')
                        after_run.append(after_t)

                        # Insert after the current run element
                        run._element.addnext(after_run)
                        run._element.addnext(kw_run)
                    break  # one keyword per run is enough

    # Body paragraphs
    _process_paragraphs(doc.paragraphs)

    # Tables (including nested)
    def _process_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    _process_paragraphs(cell.paragraphs)
                    _process_tables(cell.tables)

    _process_tables(doc.tables)

    # Headers & footers
    for section in doc.sections:
        for part in (section.header, section.footer,
                     section.first_page_header, section.first_page_footer,
                     section.even_page_header, section.even_page_footer):
            if part is None:
                continue
            _process_paragraphs(part.paragraphs)
            _process_tables(part.tables)


def _reset_run_colors(runs):
    """Set every run's font colour to black, removing template-red."""
    for r in runs:
        if not r.text:
            continue
        try:
            # python-docx way
            r.font.color.rgb = RGBColor(0, 0, 0)
        except Exception:
            # Fallback: direct XML manipulation
            try:
                rpr = r._element.find(qn('w:rPr'))
                if rpr is not None:
                    color_elem = rpr.find(qn('w:color'))
                    if color_elem is not None:
                        color_elem.set(qn('w:val'), '000000')
            except Exception:
                pass


# ── CLI ─────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    import json
    import sys

    if len(sys.argv) < 2:
        print('Usage: python tools/docx_template_filler.py <ui_json_file> [template_name]')
        print('  Fills template with values from JSON and saves to outputs/')
        sys.exit(1)

    json_path = sys.argv[1]
    template = sys.argv[2] if len(sys.argv) > 2 else 'AnlasmaBelgesi-#Dolu_v1.docx'

    with open(json_path, encoding='utf-8') as f:
        ui_data = json.load(f)

    out_name = generate_output_filename(ui_data)
    out_path = Path('outputs') / out_name

    result = fill_template(ui_data, template_name=template, output_path=str(out_path))
    print(f'Generated: {out_path} ({len(result):,} bytes)')

    # Show replacement summary
    reps = _build_replacements(ui_data)
    filled = sum(1 for v in reps.values() if v)
    print(f'Placeholders: {len(reps)} total, {filled} filled')
