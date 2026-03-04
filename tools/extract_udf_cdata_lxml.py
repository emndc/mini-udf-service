#!/usr/bin/env python3
"""Extract CDATA from content.xml inside a UDF using lxml (with fallback).

Usage: run from project root. Prints detected encoding and the CDATA content.
"""
import zipfile
import sys
import re
import io
from collections import OrderedDict

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None

def extract_cdata_with_lxml(udf_path):
    try:
        from lxml import etree
    except Exception:
        etree = None

    with zipfile.ZipFile(udf_path, 'r') as z:
        if 'content.xml' not in z.namelist():
            raise FileNotFoundError('content.xml not in ' + udf_path)
        raw = z.read('content.xml')

    # first attempt: parse with lxml and get inner text of <content>
    if etree is not None:
        try:
            parser = etree.XMLParser(recover=True, encoding='utf-8')
            root = etree.fromstring(raw, parser=parser)
            content_el = root.find('.//content')
            if content_el is not None:
                # lxml returns text without CDATA wrapper; return raw text
                return 'lxml', content_el.text or '', root
        except Exception:
            pass

    # fallback: try to extract CDATA via regex from raw bytes
    m = re.search(b'<!\[CDATA\[(.*?)\]\]>', raw, flags=re.DOTALL)
    if m:
        # try to decode as utf-8, then cp1254, then latin1
        candidates = ['utf-8', 'cp1254', 'iso-8859-9', 'latin1']
        for enc in candidates:
            try:
                text = m.group(1).decode(enc)
                return enc, text, None
            except Exception:
                continue
        # last resort: decode with errors replaced
        return 'bytes(replaced)', m.group(1).decode('utf-8', errors='replace')

    # final fallback: try a generic parse with python's minidom
    try:
        from xml.dom import minidom
        doc = minidom.parseString(raw.decode('utf-8', errors='replace'))
        elems = doc.getElementsByTagName('content')
        if elems and elems[0].firstChild:
            return 'minidom', elems[0].firstChild.nodeValue, None
    except Exception:
        pass

    # nothing found
    return None, '', None


def create_docx_from_udf_lxml(udf_path, out_docx_path):
    """Parse UDF with lxml, extract fonts used and create a DOCX preserving runs and images."""
    try:
        from lxml import etree
    except Exception:
        etree = None

    enc, text, root = extract_cdata_with_lxml(udf_path)
    # Read raw content.xml bytes to allow byte-aware CDATA extraction
    with zipfile.ZipFile(udf_path, 'r') as z:
        raw_bytes = z.read('content.xml')

    # Extract inner CDATA bytes if present
    m_bytes = re.search(b'<!\[CDATA\[(.*?)\]\]>', raw_bytes, flags=re.DOTALL)
    decoded_inner = None
    chosen_enc = None
    if m_bytes:
        inner_bytes = m_bytes.group(1)
        # try candidate decodings and pick the one with most Turkish letters and fewest replacement chars
        candidates = ['utf-8', 'cp1254', 'iso-8859-9', 'latin1']
        best = (None, -10**9)
        turkish_letters = set('çğıöşüÇĞİÖŞÜ')
        for ce in candidates:
            try:
                s = inner_bytes.decode(ce)
            except Exception:
                continue
            rep = s.count('\ufffd') + s.count('�')
            tur = sum(s.count(ch) for ch in turkish_letters)
            score = tur - rep*5
            if score > best[1]:
                best = (s, score)
                chosen_enc = ce
        if best[0] is not None:
            decoded_inner = best[0]
        else:
            decoded_inner = inner_bytes.decode('utf-8', errors='replace')
            chosen_enc = 'utf-8(replace)'

    if root is None:
        # try parse raw for elements
        with zipfile.ZipFile(udf_path, 'r') as z:
            raw = z.read('content.xml')
        if etree is None:
            raise RuntimeError('lxml not available and fallback parsing not implemented')
        parser = etree.XMLParser(recover=True, encoding='utf-8')
        root = etree.fromstring(raw, parser=parser)

    # collect font families
    fonts = OrderedDict()
    elements = root.find('.//elements')
    document = Document() if Document is not None else None

    # If we didn't find a CDATA inner string above, fall back to provided text
    if decoded_inner is None:
        decoded_inner = text or ''

    # parse styles as defaults
    style_defaults = {}
    styles_elem = root.find('.//styles')
    if styles_elem is not None:
        for s in styles_elem.findall('style'):
            name = s.get('name') or ''
            fam = s.get('family')
            size = s.get('size')
            bold = s.get('bold') == 'true'
            italic = s.get('italic') == 'true'
            style_defaults[name] = {
                'family': fam,
                'size': size,
                'bold': bold,
                'italic': italic
            }

    if elements is None:
        # fallback: single paragraph with text
        if document is None:
            return list(fonts.keys()), False, 'python-docx not installed'
        document.add_paragraph(text)
        document.save(out_docx_path)
        return list(fonts.keys()), True, out_docx_path

    for para in elements.findall('paragraph'):
        if document is None:
            # just collect fonts
            for c in para.findall('content'):
                fam = (c.get('family') or '').strip()
                if fam:
                    fonts[fam] = fonts.get(fam, 0) + 1
            continue

        p = document.add_paragraph()
        # determine paragraph style if present
        para_style_name = (para.get('style') or para.get('styleName') or para.get('name') or 'default')
        para_style = style_defaults.get(para_style_name, style_defaults.get('default', {}))
        # iterate child elements in order
        for child in list(para):
            if child.tag == 'content':
                start = int(child.get('startOffset', '0'))
                length = int(child.get('length', '0'))
                # Slice from the decoded_inner (no normalization yet) so offsets match original positions
                run_text_raw = decoded_inner[start:start+length]
                # Normalize per-run to NFC for correct rendering, but preserve slicing
                try:
                    import unicodedata
                    run_text = unicodedata.normalize('NFC', run_text_raw)
                except Exception:
                    run_text = run_text_raw
                run = p.add_run(run_text)
                # resolve attributes: prefer explicit child attrs, then paragraph style, then defaults
                fam = (child.get('family') or para_style.get('family') or '')
                # default fallback: Times New Roman
                if not fam:
                    fam = 'Times New Roman'
                if fam:
                    fonts[fam] = fonts.get(fam, 0) + 1
                    try:
                        run.font.name = fam
                    except Exception:
                        pass
                size_attr = child.get('size') or para_style.get('size')
                # default to 12pt if not specified
                if not size_attr:
                    size_attr = '12'
                if size_attr:
                    try:
                        run.font.size = Pt(float(size_attr))
                    except Exception:
                        pass
                # bold/italic
                bold_attr = child.get('bold') == 'true' if child.get('bold') is not None else para_style.get('bold')
                italic_attr = child.get('italic') == 'true' if child.get('italic') is not None else para_style.get('italic')
                if bold_attr:
                    run.bold = True
                if italic_attr:
                    run.italic = True
            elif child.tag == 'image':
                img_b64 = child.get('imageData') or (child.text or '').strip()
                if not img_b64:
                    continue
                try:
                    import base64
                    img_bytes = base64.b64decode(img_b64)
                    run = p.add_run()
                    try:
                        run.add_picture(io.BytesIO(img_bytes))
                    except Exception:
                        try:
                            document.add_picture(io.BytesIO(img_bytes))
                        except Exception:
                            pass
                except Exception:
                    pass

    if document is not None:
        # Ensure full CDATA text is also present in the document (fallback)
        try:
            # Append the decoded_inner (normalized) as a final paragraph to guarantee nothing is omitted
            if decoded_inner:
                try:
                    import unicodedata
                    full_norm = unicodedata.normalize('NFC', decoded_inner)
                except Exception:
                    full_norm = decoded_inner
                document.add_paragraph(full_norm)
        except Exception:
            pass
        document.save(out_docx_path)

    return list(fonts.keys()), True, out_docx_path


def write_cdata_only_docx(udf_path, out_docx_path):
    """Write only the CDATA decoded text into a simple DOCX (one paragraph per original newline)."""
    # extract CDATA bytes
    with zipfile.ZipFile(udf_path, 'r') as z:
        if 'content.xml' not in z.namelist():
            raise FileNotFoundError('content.xml not found')
        raw = z.read('content.xml')

    m = re.search(b'<!\[CDATA\[(.*?)\]\]>', raw, flags=re.DOTALL)
    if not m:
        # fallback to lxml extractor
        enc, text, _ = extract_cdata_with_lxml(udf_path)
    else:
        inner = m.group(1)
        # try decodings
        candidates = ['utf-8', 'cp1254', 'iso-8859-9', 'latin1']
        decoded = None
        for ce in candidates:
            try:
                s = inner.decode(ce)
                # prefer one with Turkish letters
                decoded = s
                break
            except Exception:
                continue
        if decoded is None:
            decoded = inner.decode('utf-8', errors='replace')
        enc = 'detected'  
        text = decoded

    # normalize
    try:
        import unicodedata
        text = unicodedata.normalize('NFC', text)
    except Exception:
        pass

    if Document is None:
        raise RuntimeError('python-docx not installed')

    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(out_docx_path)
    return True, out_docx_path

if __name__ == '__main__':
    path = sys.argv[1] if len(sys.argv) > 1 else 'C:\\Users\\emndc\\PycharmProjects\\flask-udf-converter\\uploads\\ornek-1.udf'
    try:
        enc, text, _ = extract_cdata_with_lxml(path)
        print('detected_by:', enc)
        print('--- CDATA START ---')
        print(text)
        print('--- CDATA END ---')
    except Exception as e:
        print('ERROR:', e)
        sys.exit(2)
